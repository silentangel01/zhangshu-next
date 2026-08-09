from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "docs" / "test-reports"
ASSET_DIR = REPORT_DIR / "assets"
OUTPUT_PATH = REPORT_DIR / "2026-07-18-material-linkage-validation.docx"

SKILL_ROOT = Path(
    r"C:\Users\JunLing\.codex\plugins\cache\openai-primary-runtime"
    r"\documents\26.715.12143\skills\documents"
)
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
PALE_BLUE = "EDF4FB"
PALE_GREEN = "EAF5EE"
GREEN = "216E39"
PALE_GOLD = "FFF6DD"
GOLD = "7A5A00"
PALE_RED = "FDECEC"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111111"


TEST_CASES = [
    (
        "ML-01",
        "服务基线",
        "启动 FastAPI 与 Vite；访问后端健康接口和前端项目页。",
        "两端均在有限时间内返回 HTTP 200。",
        "通过",
    ),
    (
        "ML-02",
        "完整联动总览",
        "打开 TEST / 章节一并进入“联动”。",
        "显示 6/6；六类卡片均有数量和摘要。",
        "通过",
    ),
    (
        "ML-03",
        "数量反向核对",
        "将总览数量与章节 API 数据逐类核对。",
        "细纲2、人物1、设定1、伏笔1、时间1完全一致。",
        "通过",
    ),
    (
        "ML-04",
        "普通分类详情",
        "依次打开细纲、人物、设定和伏笔。",
        "名称、状态、摘要与总览一致；没有跨章资料。",
        "通过",
    ),
    (
        "ML-05",
        "时间线详情",
        "打开“时间”，核对事件、前后节点和连接。",
        "显示事件1、前后节点1、相关连接2。",
        "通过",
    ),
    (
        "ML-06",
        "关系图详情",
        "打开“关系”，核对核心节点、直接关系和关联节点。",
        "显示核心节点1、直接关系2、关联节点2。",
        "通过",
    ),
    (
        "ML-07",
        "页签刷新恢复",
        "停留在关系页签后刷新写作页。",
        "刷新后仍恢复关系详情，数据重新加载正确。",
        "通过",
    ),
    (
        "ML-08",
        "章节数据隔离",
        "从完整关联章节切换到“测试章节拖拽”。",
        "关系节点、直接关系、关联节点均变为0，无旧数据残留。",
        "通过",
    ),
    (
        "ML-09",
        "快速切换竞态",
        "连续快速点击章节2和章节一。",
        "最终正文和关系数据均属于章节一，迟到请求不覆盖新状态。",
        "通过",
    ),
    (
        "ML-10",
        "部分关联总览",
        "打开雾港纪事 / 第一章 雾港来信。",
        "显示1/6；人物3；其余五类为“尚未关联”。",
        "通过",
    ),
    (
        "ML-11",
        "缺失关联空状态",
        "从细纲0卡片进入详情，并核对人物详情。",
        "细纲显示明确空状态；人物显示3项，不误显其他资料。",
        "通过",
    ),
    (
        "ML-12",
        "运行质量",
        "读取浏览器控制台，并检查资料联动接口响应。",
        "控制台 error 为0；目标资料接口均成功返回。",
        "通过",
    ),
    (
        "ML-13",
        "组件回归测试",
        "运行 chapter-context-summary.spec.ts。",
        "完整、空状态、切换事件和竞态共4条断言通过。",
        "通过",
    ),
    (
        "ML-14",
        "全量回归与构建",
        "运行完整前端单元测试和生产构建。",
        "18个文件、265条测试通过；类型检查与构建成功。",
        "通过",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges: dict[str, str | int]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, settings in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in settings.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color: str, size: int = 18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    east_asia: str = "Microsoft YaHei",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color: str, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=9, color=MUTED)


def add_metadata_paragraph(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.1
    label_run = paragraph.add_run(f"{label}：")
    set_run_font(label_run, size=10.5, color=BLACK, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=10.5, color=BLACK)


def add_callout(doc: Document, label: str, text: str, *, fill: str, border: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(paragraph, fill)
    set_paragraph_left_border(paragraph, border)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=11, color=border, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=BLACK)


def add_section_text(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=11, color=BLACK, bold=True)
        body = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(body, size=11, color=BLACK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=BLACK)


def format_table_text(table, *, body_size: float = 9.2) -> None:
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                if cell_index in (0, len(row.cells) - 1):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9.2 if row_index == 0 else body_size,
                        color=WHITE if row_index == 0 else BLACK,
                        bold=row_index == 0,
                    )
            if row_index == 0:
                set_cell_shading(cell, BLUE)
                set_cell_border(
                    cell,
                    top={"val": "single", "sz": 6, "color": BLUE},
                    bottom={"val": "single", "sz": 6, "color": BLUE},
                    start={"val": "single", "sz": 6, "color": BLUE},
                    end={"val": "single", "sz": 6, "color": BLUE},
                )
            else:
                if row_index % 2 == 0:
                    set_cell_shading(cell, "FAFBFC")
                border = {"val": "single", "sz": 4, "color": "D6DCE4"}
                set_cell_border(cell, top=border, bottom=border, start=border, end=border)
                if cell_index == len(row.cells) - 1:
                    for run in cell.paragraphs[0].runs:
                        set_run_font(run, size=9.2, color=GREEN, bold=True)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.45))

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(0)
    caption_paragraph.paragraph_format.space_after = Pt(10)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=MUTED, italic=True)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    h1 = styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run("章枢 0.9.3  ·  资料联动质量验证")
    set_run_font(header_run, size=9, color=MUTED, bold=True)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    prefix = footer_paragraph.add_run("内部测试报告  ·  第 ")
    set_run_font(prefix, size=9, color=MUTED)
    add_page_field(footer_paragraph)
    suffix = footer_paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("章枢资料联动功能验证测试报告")
    set_run_font(title_run, size=23, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run("版本 0.9.3  |  完整关联、部分关联、竞态隔离与自动化回归")
    set_run_font(subtitle_run, size=12, color=MUTED)

    add_metadata_paragraph(doc, "测试日期", "2026年7月18日")
    add_metadata_paragraph(doc, "执行环境", "Windows；FastAPI 127.0.0.1:8000；Vite 127.0.0.1:5180")
    add_metadata_paragraph(doc, "测试数据", "TEST / 章节一（6/6完整关联）与 雾港纪事 / 第一章 雾港来信（1/6部分关联）")
    add_metadata_paragraph(doc, "数据保护", "只读验证既有小说和资料；未修改用户正文、关联关系或版本历史")
    add_metadata_paragraph(doc, "执行结论", "通过，建议带改进进入下一阶段")

    add_callout(
        doc,
        "总体验收",
        "14项验证全部通过；资料数量、分类详情、空状态、跨章节隔离、刷新恢复和快速切换均符合预期。"
        "新增4条组件测试，完整前端回归265条全部通过。未发现资料联动功能性阻断缺陷。",
        fill=PALE_GREEN,
        border=GREEN,
    )

    doc.add_heading("1. 测试目标与方法", level=1)
    add_section_text(
        doc,
        "本轮在上一轮“可查看联动结果”的基础上，重点验证三件事：总览是否与真实接口数据一致；"
        "章节和页签快速切换时是否出现旧数据覆盖；完整、部分和空关联状态是否都能稳定呈现。"
    )
    add_section_text(
        doc,
        "验证方式由浏览器手工交互、REST API反向核对、浏览器控制台检查、组件自动化测试、"
        "完整前端回归和生产构建共同组成。截图仅包含章枢浏览器页面。"
    )

    doc.add_heading("2. 测试结果总览", level=1)
    add_callout(
        doc,
        "结果",
        "通过14项，失败0项，阻断0项。另记录2项高优先级改进、2项中优先级改进和2项环境/构建观察。",
        fill=PALE_BLUE,
        border=BLUE,
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["编号", "验证点", "操作与数据", "期望结果", "状态"]
    for index, text in enumerate(headers):
        table.rows[0].cells[index].text = text
    for case in TEST_CASES:
        cells = table.add_row().cells
        for index, text in enumerate(case):
            cells[index].text = text
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(
        table,
        [720, 1700, 3250, 2690, 1000],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120},
    )
    format_table_text(table, body_size=8.9)

    doc.add_heading("3. 关键验证证据", level=1)
    add_figure(
        doc,
        ASSET_DIR / "2026-07-18-material-linkage-complete.png",
        "图1  TEST / 章节一：六类资料完整关联总览（6/6）",
    )
    add_figure(
        doc,
        ASSET_DIR / "2026-07-18-material-linkage-relation-detail.png",
        "图2  TEST / 章节一：关系详情显示核心节点1、直接关系2（顶部同步失败为既有云会话失效，与本地联动无关）",
    )

    doc.add_heading("4. 自动化与构建结果", level=1)
    doc.add_heading("4.1 新增组件回归测试", level=2)
    add_section_text(
        doc,
        "新增 frontend/src/__tests__/chapter-context-summary.spec.ts，覆盖六类完整关联汇总、"
        "缺失关联提示与分类事件、快速切换时忽略迟到请求、接口失败统一错误状态。"
    )
    add_callout(
        doc,
        "执行结果",
        "1个测试文件、4条测试全部通过；测试耗时2.36秒。",
        fill=PALE_GREEN,
        border=GREEN,
    )
    doc.add_heading("4.2 全量回归与构建", level=2)
    add_section_text(
        doc,
        "完整前端单元测试：18个测试文件、265条测试全部通过。生产构建与 vue-tsc 类型检查成功。"
    )
    add_callout(
        doc,
        "构建观察",
        "主 JavaScript chunk 为639.96 kB，超过500 kB提示阈值。该提示不影响本轮正确性，"
        "但应在后续性能任务中通过页面级动态导入和代码分包处理。",
        fill=PALE_GOLD,
        border=GOLD,
    )

    doc.add_heading("5. 发现与风险", level=1)
    doc.add_heading("F-01 · P1：页签切换重复全量加载九类接口", level=2)
    add_section_text(
        doc,
        "当前组件把 projectId、chapterId 和 kind 一起作为刷新条件。用户从“人物”切到“设定”时，"
        "仍会重新请求细纲、人物、设定、伏笔、时间线轨道、时间线连接、关系节点和关系边等全部数据。"
        "小数据下没有出现错误，但长篇项目会产生不必要的数据库与渲染压力。"
    )
    add_callout(
        doc,
        "判断",
        "这是本轮最应优先处理的技术问题。它没有造成当前测试失败，却会直接影响资料规模增长后的写作流畅度。",
        fill=PALE_GOLD,
        border=GOLD,
    )

    doc.add_heading("F-02 · P1：联动页仍是查看入口，不能直接维护", level=2)
    add_section_text(
        doc,
        "缺失卡片可以进入空详情，但作者仍需跳转完整资料库才能新建或绑定。联动功能完成了“看见缺口”，"
        "尚未完成“在写作现场修复缺口”的操作闭环。"
    )

    doc.add_heading("F-03 · P2：关系卡片数量语义容易误解", level=2)
    add_section_text(
        doc,
        "联动总览的“关系1”表示本章匹配到1个核心节点；关系详情同时显示2条直接关系和2个关联节点。"
        "用户可能把总览数字理解为关系数量。建议改成“核心节点1 / 直接关系2”，或直接以关系条数作为主指标。"
    )

    doc.add_heading("F-04 · P2：浏览器回归仍依赖本机既有测试数据", level=2)
    add_section_text(
        doc,
        "本轮新增了组件测试，但完整联动、部分联动和空状态的端到端验证仍依赖本机 TEST 与雾港纪事数据。"
        "数据变化后预期值可能漂移，应建立可重复生成的测试夹具。"
    )

    doc.add_heading("环境观察", level=2)
    add_section_text(
        doc,
        "TEST 项目的云登录已过期，页面会显示“同步失败”，后端自动同步请求返回400；"
        "雾港纪事未启用云同步。两者均未影响本地资料读取、章节切换或本轮联动结果。"
    )
    add_section_text(
        doc,
        "浏览器部分关联页面的额外截图调用连续两次超时，因此报告使用完整联动和关系详情两张成功截图；"
        "部分关联状态已通过DOM、章节API与人物详情三种方式完成验证。"
    )

    doc.add_heading("6. 下一步改进意见", level=1)
    doc.add_heading("P0 · 优化联动数据加载与缓存", level=2)
    add_section_text(
        doc,
        "短期先让资料刷新只响应 projectId / chapterId 变化，页签切换复用已经加载的数据。"
        "中期增加 GET /api/chapters/{chapter_id}/context-summary 聚合接口，总览一次返回六类数量、名称摘要和"
        "关系/时间线概要；详细内容按需加载，并在关联变更后显式失效缓存。"
    )
    add_section_text(
        doc,
        "验收建议：同一章节内连续切换页签不再触发九接口全量请求；快速切换20次不出现跨章数据；"
        "本地数据下联动总览P95加载时间不高于300毫秒。"
    )

    doc.add_heading("P1 · 在写作侧栏完成直接绑定", level=2)
    add_section_text(
        doc,
        "每张卡片提供“绑定已有”“新建并绑定”“解除绑定”和“打开资料库”。缺失提示的“查看”应升级为"
        "可搜索、可多选的轻量弹窗，使作者不离开正文即可完成资料维护。"
    )
    add_section_text(
        doc,
        "验收建议：从发现缺失到完成绑定不超过3次操作；提交后总览数量立即更新；失败时保留用户选择并可重试。"
    )

    doc.add_heading("P2 · 明确统计语义与来源", level=2)
    add_section_text(
        doc,
        "关系卡片同时展示核心节点数和直接关系数；时间卡片区分本章事件和跨事件连接。所有摘要支持跳回"
        "来源资料，后续AI建议必须标明章节和原文证据，并经作者确认后写回。"
    )

    doc.add_heading("P3 · 建立可重复的端到端测试夹具", level=2)
    add_section_text(
        doc,
        "提供专用测试数据库或初始化脚本，固定生成完整关联、部分关联、空关联、隐藏节点、跨轴连接、"
        "接口迟到和接口失败场景。将本轮14项用例中的关键路径纳入CI或发布前验收。"
    )

    doc.add_heading("7. 最终结论", level=1)
    add_callout(
        doc,
        "结论",
        "资料联动功能在当前0.9.3数据规模下正确、稳定，可继续使用。下一版本应先解决重复全量加载，"
        "随后把写作页从“查看关联”提升为“直接维护关联”。这两项完成后，再推进可审阅的动态故事圣经，"
        "能够获得更可靠的性能基础和更完整的作者工作流。",
        fill=PALE_GREEN,
        border=GREEN,
    )

    doc.add_heading("附录：验证命令", level=1)
    commands = [
        "npm.cmd run test:unit -- src/__tests__/chapter-context-summary.spec.ts --run",
        "npm.cmd run test:unit -- --run",
        "npm.cmd run build",
    ]
    for command in commands:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.right_indent = Inches(0.2)
        paragraph.paragraph_format.space_after = Pt(4)
        set_paragraph_shading(paragraph, LIGHT_GRAY)
        run = paragraph.add_run(command)
        set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9.2, color=NAVY)

    return doc


def main() -> None:
    doc = build_document()
    doc.core_properties.title = "章枢资料联动功能验证测试报告"
    doc.core_properties.subject = "章枢 0.9.3 资料联动验证"
    doc.core_properties.author = "章枢测试"
    doc.core_properties.keywords = "章枢, 资料联动, 测试报告, 0.9.3"
    doc.save(OUTPUT_PATH)
    verify_document(OUTPUT_PATH)
    print(OUTPUT_PATH)


def verify_document(path: Path) -> None:
    doc = Document(path)
    assert len(doc.sections) == 1
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11.0
    assert all(
        round(value.inches, 2) == 1.0
        for value in (
            section.top_margin,
            section.right_margin,
            section.bottom_margin,
            section.left_margin,
        )
    )
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 15
    assert len(doc.tables[0].columns) == 5
    assert len(doc.inline_shapes) == 2

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    for required in (
        "章枢资料联动功能验证测试报告",
        "14项验证全部通过",
        "18个测试文件、265条测试全部通过",
        "F-01 · P1：页签切换重复全量加载九类接口",
        "P0 · 优化联动数据加载与缓存",
        "最终结论",
    ):
        assert required in text, required
    for forbidden in ("turn0", "cite", "TODO", "PLACEHOLDER"):
        assert forbidden not in text, forbidden

    headings = [
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.style.name.startswith("Heading")
    ]
    print(
        {
            "paragraphs": len(doc.paragraphs),
            "headings": len(headings),
            "tables": len(doc.tables),
            "table_rows": len(doc.tables[0].rows),
            "images": len(doc.inline_shapes),
            "page_inches": (
                round(section.page_width.inches, 2),
                round(section.page_height.inches, 2),
            ),
            "margins_inches": (
                round(section.top_margin.inches, 2),
                round(section.right_margin.inches, 2),
                round(section.bottom_margin.inches, 2),
                round(section.left_margin.inches, 2),
            ),
            "required_text_checks": "passed",
            "forbidden_token_checks": "passed",
        }
    )


if __name__ == "__main__":
    main()
