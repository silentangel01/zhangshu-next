from __future__ import annotations

from io import BytesIO

from docx import Document

from app.services.export_service import ManuscriptDocument


def render_manuscript_docx(document: ManuscriptDocument) -> bytes:
    doc = Document()

    doc.add_heading(document.project_title, level=0)

    for volume in document.volumes:
        if volume.title:
            doc.add_heading(volume.title, level=1)

        for chapter in volume.chapters:
            doc.add_heading(chapter.title, level=2)

            if chapter.content and chapter.content.strip():
                for paragraph in chapter.content.split("\n"):
                    text = paragraph.rstrip()
                    if text:
                        doc.add_paragraph(text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
